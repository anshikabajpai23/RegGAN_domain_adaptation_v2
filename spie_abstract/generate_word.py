"""
Generate SPIE abstract as a formatted Word (.docx) document.
Run: python3 generate_word.py
Output: spie_abstract.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (SPIE letter-size: 0.88" L/R, 1.0" top, 1.25" bottom) ─────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.25)
    section.left_margin   = Inches(0.88)
    section.right_margin  = Inches(0.88)

# ── Default body font ────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def para(text="", bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, indent=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    if text:
        run = p.add_run(text)
        run.font.name   = 'Times New Roman'
        run.font.size   = Pt(size)
        run.bold        = bold
        run.italic      = italic
    return p

def mixed_para(parts, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
               space_after=6, indent=False, size=12):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name   = 'Times New Roman'
        run.font.size   = Pt(size)
        run.bold        = bold
        run.italic      = italic
    return p

def section_heading(number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"{number}. {title.upper()}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    return p

def table_row(tbl, cells_data, bold=False, shade=None):
    row = tbl.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.bold = bold
    return row

# ============================================================
# TITLE
# ============================================================
para("Cross-Modality Transfer Learning for Knee Meniscus Segmentation",
     bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=0, space_after=0)
para("via Registration-Constrained Unpaired Image Translation",
     bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=0, space_after=10)

# AUTHORS
para("Anshika Bajpaiᵃ, Chia-Ying James Linᵃ, Madilyn Feikᶜ, Abhay Sistaᶜ, Mounica Chiduralaᶜ, "
     "Ashley Ellenbergerᶜ, Bryan Saltzmanᶜ, and Rakesh Shiradkarᵇ",
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=4)

# AFFILIATIONS
para("ᵃIndiana University Bloomington, Bloomington, Indiana; "
     "ᵇDepartment of Biomedical Engineering and Informatics, "
     "Indiana University Indianapolis, Indianapolis, Indiana; "
     "ᶜIU Health, Indianapolis, Indiana",
     italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=0, space_after=14)

# ============================================================
# ABSTRACT
# ============================================================
para("ABSTRACT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=6, space_after=6)

para(
    "Automated meniscus segmentation in proton density (PD)-weighted knee MRI is "
    "hindered by the scarcity of labeled training data. Dual-echo steady-state (DESS) "
    "MRI offers rich publicly available annotations, but the fundamental contrast "
    "difference between DESS and PD (fluid appears dark in DESS and bright in "
    "PD), preventing direct model transfer. We propose a cross-modality domain adaptation "
    "pipeline that uses a registration-constrained generative adversarial network "
    "(RegGAN) to translate DESS volumes to synthetic PD images while explicitly "
    "penalizing anatomical deformation. The translation preserves meniscus anatomy "
    "(mean displacement 0.055 pixels, 0% topology violations), enabling DESS "
    "segmentation masks to supervise a 2.5D U-Net in the PD domain. A pre-trained "
    "DESS segmentation model applied directly to real PD achieves near-zero meniscus "
    "Dice (0.00). Fine-tuning on synthetic PD images raises Dice to 0.69 on a "
    "held-out cohort of 17 expert-annotated real PD patients. Further fine-tuning "
    "with 15 manually annotated real PD cases achieves a mean Dice of "
    "0.78, demonstrating that anatomy-preserving synthetic data substantially reduces "
    "annotation burden for cross-modality segmentation.",
    space_before=0, space_after=8, indent=False
)

# KEYWORDS
mixed_para([
    ("Keywords: ", True, False),
    ("Knee MRI, Meniscus segmentation, Domain adaptation, Generative adversarial "
     "networks, Unpaired image translation, Transfer learning, Cross-modality", False, False),
], space_before=0, space_after=14)

# ============================================================
# 1. DESCRIPTION OF PURPOSE
# ============================================================
section_heading(1, "Description of Purpose")

para(
    "The meniscus is a critical fibrocartilaginous structure responsible for load "
    "distribution, shock absorption, and joint stability in the knee. Accurate automated "
    "meniscus segmentation in MRI supports quantitative grading of cartilage injury, "
    "longitudinal monitoring of osteoarthritis progression, and pre-surgical planning. "
    "Proton density (PD)-weighted MRI is the dominant clinical protocol for knee "
    "evaluation, offering high soft-tissue contrast and short acquisition times. However, "
    "constructing supervised segmentation models for PD MRI requires expert voxel-level "
    "annotation, which is expensive and time-intensive at scale.",
    space_after=6, indent=True
)

para(
    "Dual-echo steady-state (DESS) MRI provides complementary T2 relaxometry and has "
    "been distributed alongside dense expert segmentation labels in publicly available "
    "datasets such as SKM-TEA [1]. The tissue contrast between the two sequences differs "
    "substantially: fluid and cartilage appear dark on DESS but bright on PD, with "
    "reversed contrast at tissue boundaries, making direct transfer of a DESS-trained "
    "model to PD images unsuccessful. Prior unpaired image-to-image translation methods "
    "such as CycleGAN [2] have achieved modality-level appearance transfer but produce "
    "uncontrolled spatial deformations, distorting meniscus anatomy and rendering "
    "associated segmentation masks unreliable for downstream training.",
    space_after=6, indent=True
)

para(
    "This work addresses this limitation by introducing a registration-constrained "
    "generative pipeline (RegGAN) that enforces minimum geometric deformation during "
    "DESS-to-PD translation. We evaluate whether synthetic PD images generated under "
    "anatomy-preserving constraints are sufficient to transfer meniscus segmentation "
    "knowledge from the DESS domain to real PD MRI, and whether augmenting synthetic "
    "training data with a small cohort of manually annotated real PD cases further "
    "improves segmentation on unseen clinical data.",
    space_after=12, indent=True
)

# ============================================================
# 2. METHODS
# ============================================================
section_heading(2, "Methods")

mixed_para([
    ("Domain adaptation. ", True, False),
    ("RegGAN extends the CycleGAN framework [2] with a lightweight differentiable "
     "registration network R, implemented as a 2D convolutional U-Net inspired by "
     "VoxelMorph [3]. The generator G", False, False),
    ("AB", False, False),
    (" translates DESS slices to synthetic PD; R then takes the concatenation of the "
     "synthetic PD and a real PD image as input and predicts a dense 2D displacement "
     "field. Three auxiliary loss terms constrain the predicted field: a registration "
     "similarity loss (λ = 5.0) enforcing pixel-level alignment between "
     "the warped synthetic PD and real PD, a smoothness penalty (λ = 10.0) "
     "on spatial gradients of the displacement field, and a magnitude penalty "
     "(λ = 5.0) directly penalizing displacement amplitude. Generators "
     "and discriminators follow a ResNet-based [4] architecture with 9 residual blocks "
     "(nᵍgf = 48) and PatchGAN [5] discriminators trained using the "
     "least-squares GAN objective [6]. Training used 69 DESS volumes from SKM-TEA [1] "
     "and 69 PD-weighted volumes from an institutional cohort (IU dataset), extracting "
     "approximately 11,000 DESS and 2,200 PD sagittal slices as unpaired training data "
     "on an NVIDIA A100 40 GB GPU.", False, False),
], space_after=6, indent=True)

mixed_para([
    ("Segmentation and fine-tuning. ", True, False),
    ("A publicly available 2.5D U-Net with ResNet34 encoder pre-trained on DESS for "
     "multi-class knee segmentation [7] serves as the source-domain baseline. The "
     "original 5-class head is replaced by a 3-class head (background, lateral meniscus, "
     "medial meniscus) for cross-modality transfer. In the first fine-tuning configuration "
     "(", False, False),
    ("synthetic PD", False, True),
    ("), only synthetic PD slices paired with DESS meniscus labels are used. In the "
     "second configuration (", False, False),
    ("synthetic + real PD", False, True),
    ("), 15 manually annotated real PD patients are added to training "
     "using a merged binary loss that combines lateral and medial softmax probabilities "
     "against binary ground truth. Both configurations use weighted cross-entropy "
     "(background weight 0.1, meniscus weight 1.5) combined with soft Dice loss, "
     "cosine learning rate scheduling (1×10⁻⁵ to 1×10⁻⁷), "
     "and augmentation including horizontal/vertical flipping, brightness jitter (±20%), "
     "and Gaussian noise (σ = 0.02), trained for 50 epochs with early "
     "stopping (patience 10).", False, False),
], space_after=6, indent=True)

mixed_para([
    ("Evaluation. ", True, False),
    ("All segmentation models were evaluated on a held-out cohort of 17 real PD patients "
     "with expert manual meniscus annotations, distinct from the fine-tuning cohort. "
     "Performance was quantified using the binary Dice coefficient (lateral and medial "
     "meniscus combined). Domain adaptation quality was assessed with Fréchet "
     "Inception Distance (FID) [8], structural similarity index (SSIM), and Jacobian "
     "determinant analysis of the deformation field to confirm topology preservation.", False, False),
], space_after=12, indent=True)

# ============================================================
# 3. RESULTS AND EVALUATION
# ============================================================
section_heading(3, "Results and Evaluation")

para(
    "RegGAN achieved an FID of 164.4 between synthetic PD and real PD images, compared "
    "to 260.4 for raw DESS images evaluated against the same real PD distribution, a "
    "37% reduction in distributional distance. The SSIM between DESS and synthetic PD "
    "was 0.357, reflecting the expected contrast shift rather than structural failure. "
    "Deformation analysis confirmed anatomy preservation throughout: the mean Jacobian "
    "determinant of the registration field was 1.000056; the rate of Jacobian folding "
    "(det(J) < 0, indicating topology violations) was 0.0% across all test "
    "volumes; and the mean displacement at the meniscus region was 0.055 pixels.",
    space_after=6, indent=True
)

para(
    "Segmentation performance on the 17-patient held-out cohort is summarized in "
    "Table 1. The pre-trained DESS model applied without fine-tuning to real PD "
    "achieves a mean binary meniscus Dice of 0.00, confirming complete failure of "
    "zero-shot cross-modality transfer. Fine-tuning on synthetic PD images alone "
    "improves mean Dice to 0.69, demonstrating that RegGAN-generated images provide "
    "sufficient domain alignment for meaningful label transfer. The combined fine-tuning "
    "approach using synthetic PD and 15 annotated real PD patients achieves "
    "the highest performance with a mean Dice of 0.78, a relative gain of 13% over "
    "the synthetic-only baseline.",
    space_after=6, indent=True
)

para(
    "Per-patient Dice scores for the best model ranged from 0.69 to 0.84 across all "
    "17 cases, with consistent improvement over both the DESS baseline and the "
    "synthetic-only model. These results indicate that anatomy-preserving synthetic data "
    "from RegGAN effectively bridges the cross-modality gap, and that a small real-PD "
    "annotation effort provides an additional, complementary performance gain.",
    space_after=12, indent=True
)

# ============================================================
# 4. SUPPORTING IMAGES, TABLES, FIGURES
# ============================================================
section_heading(4, "Supporting Images, Tables, Figures")

p_fig1 = doc.add_paragraph()
p_fig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig1.paragraph_format.space_before = Pt(6)
p_fig1.paragraph_format.space_after  = Pt(4)
run_fig1 = p_fig1.add_run()
run_fig1.add_picture(
    "../pipeline-diagram.png",
    width=Inches(6.0)
)
para("Figure 1. Overview of the proposed cross-modality domain adaptation pipeline. "
     "RegGAN translates DESS MRI (SKM-TEA, with lateral meniscus, medial meniscus, "
     "and cartilage labels) to synthetic PD MRI under minimum-deformation constraints. "
     "A pre-trained 2.5D U-Net is progressively fine-tuned on synthetic PD and "
     "evaluated on a held-out real PD cohort.",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)

# Results table
para("Table 1. Binary meniscus Dice score on 17 held-out real PD patients.",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Header
hdr = tbl.rows[0].cells
for cell, text in zip(hdr, ["Model", "Training Data", "Mean Dice"]):
    cell.text = text
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ("Pre-trained DESS baseline",       "DESS only (no PD)",               "0.00"),
    ("Synthetic PD fine-tuning",         "Fake PD (RegGAN output)",         "0.69"),
    ("Synthetic + Real PD fine-tuning",  "Fake PD + annotated real PD",     "0.78"),
]
for row_data in rows_data:
    row = tbl.add_row()
    for cell, text in zip(row.cells, row_data):
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Dice bar chart (Figure 2)
p_fig2 = doc.add_paragraph()
p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig2.paragraph_format.space_before = Pt(10)
p_fig2.paragraph_format.space_after  = Pt(4)
run_fig2 = p_fig2.add_run()
run_fig2.add_picture(
    "../dice_abstract_v8_3models.png",
    width=Inches(6.0)
)
para("Figure 2. Per-patient binary meniscus Dice scores on all 17 held-out real PD "
     "patients (10 original + 7 extended hold-out). Dashed lines: group means. Grey: "
     "DESS baseline (no fine-tuning). Blue: synthetic PD fine-tuning. Orange: "
     "synthetic + annotated real PD fine-tuning.",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)

# Boundary overlay (Figure 3)
p_fig3 = doc.add_paragraph()
p_fig3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig3.paragraph_format.space_before = Pt(10)
p_fig3.paragraph_format.space_after  = Pt(4)
run_fig3 = p_fig3.add_run()
run_fig3.add_picture(
    "../boundary_abstract_multipatient.png",
    width=Inches(6.5)
)
para("Figure 3. Qualitative meniscus segmentation on three representative real PD patients "
     "(rows). Columns: PD image with ground truth (white), Baseline (DESS, no fine-tuning), "
     "Synthetic PD fine-tuning (pseudo PD), and Synthetic + Real PD fine-tuning (mixed). "
     "Top two rows show successful segmentation; bottom row shows a challenging case "
     "where fine-tuning partially recovers the meniscus boundary.",
     size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=12)

# ============================================================
# 5. CONCLUSIONS
# ============================================================
section_heading(5, "Conclusions")

para(
    "This study demonstrates that registration-constrained generative domain adaptation "
    "enables effective cross-modality transfer of meniscus segmentation from DESS to "
    "PD-weighted knee MRI. The proposed RegGAN pipeline achieves near-rigid anatomical "
    "preservation during translation (mean displacement 0.055 px, 0% topology "
    "violations), allowing DESS segmentation labels to be directly reused for PD model "
    "training without paired acquisitions. On a held-out cohort of 17 expert-annotated "
    "real PD patients, fine-tuning on synthetic PD alone improves meniscus Dice from "
    "0.00 to 0.69; augmenting with 15 annotated real PD patients further improves "
    "Dice to 0.78. These results establish anatomy-preserving GAN-based domain "
    "adaptation as a practical, label-efficient pathway for cross-modality MRI "
    "segmentation.",
    space_after=12, indent=True
)

# ============================================================
# 6. BREAKTHROUGH WORK
# ============================================================
section_heading(6, "Breakthrough Work")

mixed_para([
    ("This work presents a novel anatomy-preserving cross-modality adaptation framework "
     "combining a deformation-penalized registration network within an unpaired GAN "
     "training loop with a progressive transfer learning strategy for knee meniscus "
     "segmentation. By demonstrating that synthetic images generated under minimum-"
     "deformation constraints preserve the geometric fidelity needed to transfer dense "
     "segmentation labels across MRI contrasts, this study provides a reproducible and "
     "label-efficient approach applicable to other organ systems and acquisition protocol "
     "pairs where labeled source-domain data exist but target-domain annotations are "
     "scarce. ", False, False),
    ("This work has not been presented or submitted for publication or presentation "
     "elsewhere.", True, False),
], space_after=12, indent=True)

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("ACKNOWLEDGMENTS")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

para(
    "The authors gratefully acknowledge computational resources provided by Indiana "
    "University's BigRed200 high-performance computing cluster. This work was supported "
    "in part by Indiana University research funds.",
    space_after=12, indent=True
)

# ============================================================
# REFERENCES
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("REFERENCES")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

refs = [
    ("[1] Desai, A. D., Bhave, S., Hancu, I., Watkins, R. D., Gold, G. E., Hargreaves, B. A., "
     "and Sandino, C. M., \"SKM-TEA: A dataset for accelerated MRI reconstruction with dense "
     "image labels for quantitative clinical evaluation,\" in Advances in Neural Information "
     "Processing Systems, vol. 35, 2022."),
    ("[2] Zhu, J.-Y., Park, T., Isola, P., and Efros, A. A., \"Unpaired image-to-image "
     "translation using cycle-consistent adversarial networks,\" in Proc. IEEE Int. Conf. "
     "Computer Vision, 2223–2232, 2017."),
    ("[3] Balakrishnan, G., Zhao, A., Sabuncu, M. R., Guttag, J., and Dalca, A. V., "
     "\"VoxelMorph: A learning framework for deformable medical image registration,\" "
     "IEEE Trans. Medical Imaging 38(8), 1788–1800 (2019)."),
    ("[4] He, K., Zhang, X., Ren, S., and Sun, J., \"Deep residual learning for image "
     "recognition,\" in Proc. IEEE Conf. Computer Vision and Pattern Recognition, "
     "770–778, 2016."),
    ("[5] Isola, P., Zhu, J.-Y., Zhou, T., and Efros, A. A., \"Image-to-image translation "
     "with conditional adversarial networks,\" in Proc. IEEE Conf. Computer Vision and "
     "Pattern Recognition, 1125–1134, 2017."),
    ("[6] Mao, X., Li, Q., Xie, H., Lau, R. Y. K., Wang, Z., and Smolley, S. P., "
     "\"Least squares generative adversarial networks,\" in Proc. IEEE Int. Conf. "
     "Computer Vision, 2794–2802, 2017."),
    ("[7] Desai, A. D., Schmidt, A. M., Rubin, E. B., Sandino, C. M., Black, M. S., "
     "Mazzoli, V., Stevens, K. J., Boutin, R., Re, T. J., Gold, G. E., Hargreaves, B. A., "
     "and Chaudhari, A. S., \"The International Workshop on Osteoarthritis Imaging Kaggle "
     "AKOA knee MRI plane challenge results,\" Radiology: Artificial Intelligence 3(6), "
     "e200165 (2021)."),
    ("[8] Heusel, M., Ramsauer, H., Unterthiner, T., Nessler, B., and Hochreiter, S., "
     "\"GANs trained by a two time-scale update rule converge to a local Nash equilibrium,\" "
     "in Advances in Neural Information Processing Systems, vol. 30, 2017."),
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc.save("spie_abstract.docx")
print("Saved: spie_abstract.docx")
